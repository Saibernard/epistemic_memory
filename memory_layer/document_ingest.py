"""
Document Ingestion for the Memory Layer.

Extracts text from various file formats and URLs, then chunks it
into memory-sized pieces for storage and retrieval.

Supported formats:
  - PDF (.pdf)
  - Word (.docx)
  - Plain text (.txt)
  - Markdown (.md)
  - CSV (.csv)
  - JSON (.json)
  - URLs (any web page, Confluence, docs sites, etc.)

Usage:
    from memory_layer.document_ingest import DocumentIngestor

    ingestor = DocumentIngestor()

    # From file
    chunks = ingestor.extract_and_chunk("report.pdf")

    # From URL
    chunks = ingestor.extract_and_chunk_url("https://docs.example.com/api")

    for chunk in chunks:
        brain.remember(chunk["content"], tags=chunk["tags"], metadata=chunk["metadata"])
"""

import os
import io
import csv
import json
import re
from typing import List, Dict, Any, Optional, BinaryIO
from urllib.parse import urlparse


# Supported file extensions
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".csv", ".json"}

# URL fetch settings
MAX_URL_CONTENT = 500_000   # 500 KB max text from a single URL
URL_TIMEOUT = 30            # seconds

# Chunking defaults
DEFAULT_CHUNK_SIZE = 1500       # characters per chunk
DEFAULT_CHUNK_OVERLAP = 300     # overlap between chunks
MAX_DOCUMENT_SIZE = 50_000_000  # 50 MB max file size
MAX_CHUNKS_PER_DOCUMENT = 2000  # supports ~1000-page books


class DocumentIngestor:
    """
    Extracts text from documents and splits them into memory-ready chunks.

    Each chunk is returned as a dict with:
      - content: the text content to store
      - tags: auto-generated tags (filename, format, chunk index)
      - metadata: source file info, page numbers, section headers

    For PDFs: automatically uses Gemini vision when GOOGLE_API_KEY is set,
    which captures formulas, images, diagrams, and tables that text-only
    extraction misses. Falls back to pypdf text extraction otherwise.
    """

    def __init__(
        self,
        chunk_size: int = DEFAULT_CHUNK_SIZE,
        chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
        vision_model: str = "gemini-2.5-flash",
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.vision_model = vision_model

    # ─────────────────────────────────────────────
    # PUBLIC INTERFACE
    # ─────────────────────────────────────────────

    def extract_from_file(self, file_path: str) -> str:
        """Extract raw text from a file on disk."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_size = os.path.getsize(file_path)
        if file_size > MAX_DOCUMENT_SIZE:
            raise ValueError(
                f"File too large: {file_size / 1_000_000:.1f} MB "
                f"(max {MAX_DOCUMENT_SIZE / 1_000_000:.0f} MB)"
            )

        ext = os.path.splitext(file_path)[1].lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {ext}. "
                f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        with open(file_path, "rb") as f:
            return self._extract_text(f, ext, os.path.basename(file_path))

    def extract_from_bytes(
        self, data: bytes, filename: str
    ) -> str:
        """Extract raw text from in-memory bytes (e.g., file upload)."""
        if len(data) > MAX_DOCUMENT_SIZE:
            raise ValueError(
                f"File too large: {len(data) / 1_000_000:.1f} MB "
                f"(max {MAX_DOCUMENT_SIZE / 1_000_000:.0f} MB)"
            )

        ext = os.path.splitext(filename)[1].lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {ext}. "
                f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        return self._extract_text(io.BytesIO(data), ext, filename)

    def chunk_text(
        self,
        text: str,
        source_filename: str = "unknown",
        extra_tags: Optional[List[str]] = None,
        extra_metadata: Optional[Dict[str, Any]] = None,
    ) -> List[Dict[str, Any]]:
        """
        Split text into overlapping, paragraph-aware chunks.

        Returns a list of dicts, each with:
          - content: chunk text
          - tags: auto + user tags
          - metadata: source info + chunk position
        """
        if not text or not text.strip():
            return []

        ext = os.path.splitext(source_filename)[1].lower()
        base_name = os.path.splitext(os.path.basename(source_filename))[0]

        # Split into paragraphs first for semantic boundaries
        paragraphs = self._split_paragraphs(text)

        # Merge small paragraphs and split large ones to target chunk size
        raw_chunks = self._merge_and_split(paragraphs)

        if len(raw_chunks) > MAX_CHUNKS_PER_DOCUMENT:
            print(
                f"  ⚠ Document '{source_filename}' produced {len(raw_chunks)} chunks "
                f"but cap is {MAX_CHUNKS_PER_DOCUMENT}. "
                f"Last {len(raw_chunks) - MAX_CHUNKS_PER_DOCUMENT} chunks will be dropped."
            )
            raw_chunks = raw_chunks[:MAX_CHUNKS_PER_DOCUMENT]

        # Build structured chunks
        total = len(raw_chunks)
        chunks = []
        for i, chunk_text in enumerate(raw_chunks):
            chunk_text = chunk_text.strip()
            if not chunk_text:
                continue

            # Auto-detect section headers in the chunk
            section = self._detect_section_header(chunk_text)

            tags = ["document", base_name.lower().replace(" ", "_")]
            if ext:
                tags.append(ext.lstrip("."))
            if section:
                tags.append(section.lower().replace(" ", "_")[:40])
            if extra_tags:
                tags.extend(extra_tags)

            metadata = {
                "source_file": source_filename,
                "chunk_index": i,
                "total_chunks": total,
                "document_type": ext.lstrip(".") if ext else "text",
            }
            if section:
                metadata["section"] = section
            if extra_metadata:
                metadata.update(extra_metadata)

            # Prefix chunk with document context for better recall
            prefix = f"[From: {source_filename}"
            if section:
                prefix += f" | Section: {section}"
            prefix += f" | Part {i + 1}/{total}]"

            chunks.append({
                "content": f"{prefix}\n{chunk_text}",
                "tags": list(set(tags)),
                "metadata": metadata,
            })

        return chunks

    def extract_and_chunk(
        self,
        file_path: str,
        extra_tags: Optional[List[str]] = None,
        extra_metadata: Optional[Dict[str, Any]] = None,
    ) -> List[Dict[str, Any]]:
        """Extract text from file and return memory-ready chunks."""
        text = self.extract_from_file(file_path)
        return self.chunk_text(
            text,
            source_filename=os.path.basename(file_path),
            extra_tags=extra_tags,
            extra_metadata=extra_metadata,
        )

    def extract_and_chunk_bytes(
        self,
        data: bytes,
        filename: str,
        extra_tags: Optional[List[str]] = None,
        extra_metadata: Optional[Dict[str, Any]] = None,
    ) -> List[Dict[str, Any]]:
        """Extract text from in-memory bytes and return memory-ready chunks."""
        text = self.extract_from_bytes(data, filename)
        return self.chunk_text(
            text,
            source_filename=filename,
            extra_tags=extra_tags,
            extra_metadata=extra_metadata,
        )

    # ─────────────────────────────────────────────
    # URL EXTRACTION
    # ─────────────────────────────────────────────

    def extract_from_url(self, url: str) -> Dict[str, Any]:
        """
        Fetch a web page and extract its main text content.

        Uses trafilatura for clean article extraction (strips nav, ads,
        sidebars, etc.) with BeautifulSoup as a fallback.

        Returns a dict with:
          - text: the extracted text content
          - title: page title (if available)
          - url: the original URL
          - domain: the domain name
        """
        parsed = urlparse(url)
        if not parsed.scheme or not parsed.netloc:
            raise ValueError(f"Invalid URL: {url}")

        # Fetch the page
        try:
            import requests as _requests
        except ImportError:
            raise ImportError(
                "requests is required for URL ingestion. "
                "Install with: pip install requests"
            )

        headers = {
            "User-Agent": (
                "Mozilla/5.0 (compatible; MemoryLayer/1.0; "
                "+https://github.com/Saibernard/insane_memory)"
            ),
            "Accept": "text/html,application/xhtml+xml,*/*",
        }

        try:
            resp = _requests.get(
                url, headers=headers, timeout=URL_TIMEOUT, allow_redirects=True
            )
            resp.raise_for_status()
        except _requests.exceptions.Timeout:
            raise ValueError(f"URL fetch timed out after {URL_TIMEOUT}s: {url}")
        except _requests.exceptions.RequestException as e:
            raise ValueError(f"Failed to fetch URL: {e}")

        content_type = resp.headers.get("Content-Type", "")
        html = resp.text

        if not html or not html.strip():
            raise ValueError(f"Empty page content from: {url}")

        # Extract text using trafilatura (best quality)
        text = None
        title = None

        try:
            import trafilatura
            downloaded = trafilatura.extract(
                html,
                include_comments=False,
                include_tables=True,
                include_links=False,
                output_format="txt",
                favor_recall=True,
            )
            if downloaded and len(downloaded.strip()) > 50:
                text = downloaded

            # Try to get title
            metadata = trafilatura.extract_metadata(html)
            if metadata and metadata.title:
                title = metadata.title
        except ImportError:
            pass  # Fall through to BeautifulSoup
        except Exception:
            pass  # Fall through to BeautifulSoup

        # Fallback: BeautifulSoup
        if not text:
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(html, "html.parser")

                # Get title
                if not title:
                    title_tag = soup.find("title")
                    if title_tag:
                        title = title_tag.get_text(strip=True)

                # Remove script, style, nav, footer, header elements
                for tag in soup(["script", "style", "nav", "footer",
                                 "header", "aside", "noscript", "iframe"]):
                    tag.decompose()

                # Get main content area if available
                main = (
                    soup.find("main")
                    or soup.find("article")
                    or soup.find(attrs={"role": "main"})
                    or soup.find(id=re.compile(r"content|main|article", re.I))
                    or soup.find(class_=re.compile(r"content|main|article", re.I))
                    or soup.body
                    or soup
                )

                text = main.get_text(separator="\n", strip=True)
            except ImportError:
                raise ImportError(
                    "beautifulsoup4 is required for URL ingestion. "
                    "Install with: pip install beautifulsoup4"
                )

        if not text or len(text.strip()) < 20:
            raise ValueError(
                f"Could not extract meaningful text from: {url}"
            )

        # Truncate if too large
        text = text[:MAX_URL_CONTENT]
        text = self._clean_text(text)

        return {
            "text": text,
            "title": title or parsed.netloc,
            "url": url,
            "domain": parsed.netloc,
        }

    def extract_and_chunk_url(
        self,
        url: str,
        extra_tags: Optional[List[str]] = None,
        extra_metadata: Optional[Dict[str, Any]] = None,
    ) -> List[Dict[str, Any]]:
        """
        Fetch a URL, extract text, and return memory-ready chunks.

        Each chunk is tagged with the source URL and domain.
        Works with any web page: docs sites, Confluence, blog posts,
        GitHub READMEs, wikis, etc.
        """
        result = self.extract_from_url(url)
        text = result["text"]
        title = result["title"]
        domain = result["domain"]

        # Build a friendly source name
        source_name = title if title else domain

        # Combine tags
        tags = list(extra_tags or [])
        tags.extend(["url", domain.replace(".", "_")])

        # Build metadata
        meta = {
            "source_url": url,
            "domain": domain,
            "page_title": title,
            "source_type": "url",
        }
        if extra_metadata:
            meta.update(extra_metadata)

        chunks = self.chunk_text(
            text=text,
            source_filename=source_name,
            extra_tags=tags,
            extra_metadata=meta,
        )

        return chunks

    # ─────────────────────────────────────────────
    # YOUTUBE TRANSCRIPT EXTRACTION
    # ─────────────────────────────────────────────

    @staticmethod
    def is_youtube_url(url: str) -> bool:
        """Check if a URL is a YouTube video link."""
        parsed = urlparse(url)
        host = (parsed.netloc or "").lower().replace("www.", "")
        return host in ("youtube.com", "youtu.be", "m.youtube.com")

    @staticmethod
    def extract_video_id(url: str) -> Optional[str]:
        """Extract the video ID from various YouTube URL formats."""
        parsed = urlparse(url)
        host = (parsed.netloc or "").lower().replace("www.", "")

        if host == "youtu.be":
            return parsed.path.lstrip("/").split("/")[0] or None

        from urllib.parse import parse_qs
        if host in ("youtube.com", "m.youtube.com"):
            if parsed.path.startswith("/watch"):
                qs = parse_qs(parsed.query)
                return qs.get("v", [None])[0]
            if parsed.path.startswith(("/embed/", "/v/", "/shorts/")):
                return parsed.path.split("/")[2] if len(parsed.path.split("/")) > 2 else None

        vid_match = re.search(r"[?&]v=([a-zA-Z0-9_-]{11})", url)
        if vid_match:
            return vid_match.group(1)

        return None

    def extract_from_youtube(self, url: str) -> Dict[str, Any]:
        """
        Fetch the transcript (captions) of a YouTube video.

        Returns:
          - text: full transcript as readable text
          - title: video title (fetched from page)
          - video_id: the YouTube video ID
          - url: original URL
          - duration_seconds: approximate video duration
          - is_generated: whether the transcript was auto-generated
        """
        video_id = self.extract_video_id(url)
        if not video_id:
            raise ValueError(f"Could not extract video ID from URL: {url}")

        try:
            from youtube_transcript_api import YouTubeTranscriptApi
        except ImportError:
            raise ImportError(
                "youtube-transcript-api is required for YouTube support. "
                "Install with: pip install youtube-transcript-api"
            )

        ytt_api = YouTubeTranscriptApi()
        try:
            transcript = ytt_api.fetch(video_id, languages=["en", "en-US", "en-GB"])
        except Exception:
            try:
                transcript = ytt_api.fetch(video_id)
            except Exception as e:
                raise ValueError(
                    f"Could not fetch transcript for video {video_id}. "
                    f"The video may not have captions available. Error: {e}"
                )

        snippets = transcript.to_raw_data()
        if not snippets:
            raise ValueError(f"Empty transcript for video: {video_id}")

        lines = []
        for s in snippets:
            text = s.get("text", "").strip()
            if text:
                lines.append(text)

        full_text = " ".join(lines)
        full_text = re.sub(r"\s+", " ", full_text).strip()

        duration = 0
        if snippets:
            last = snippets[-1]
            duration = last.get("start", 0) + last.get("duration", 0)

        title = self._fetch_youtube_title(url, video_id)

        return {
            "text": full_text,
            "title": title,
            "video_id": video_id,
            "url": url,
            "duration_seconds": round(duration),
            "is_generated": getattr(transcript, "is_generated", True),
            "language": getattr(transcript, "language", "unknown"),
        }

    def _fetch_youtube_title(self, url: str, video_id: str) -> str:
        """Best-effort fetch of the video title from the page."""
        try:
            import requests as _requests
            resp = _requests.get(
                url,
                headers={"User-Agent": "Mozilla/5.0"},
                timeout=10,
            )
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(resp.text, "html.parser")
            title_tag = soup.find("title")
            if title_tag:
                title = title_tag.get_text(strip=True)
                title = title.replace(" - YouTube", "").strip()
                if title:
                    return title
        except Exception:
            pass
        return f"YouTube Video {video_id}"

    def extract_and_chunk_youtube(
        self,
        url: str,
        extra_tags: Optional[List[str]] = None,
        extra_metadata: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """
        Fetch a YouTube video transcript and return memory-ready chunks
        plus video metadata.
        """
        result = self.extract_from_youtube(url)
        text = result["text"]
        title = result["title"]
        video_id = result["video_id"]

        duration_min = result["duration_seconds"] / 60
        header = f"YouTube Video: {title}\n"
        header += f"Duration: {duration_min:.0f} minutes | "
        header += f"Language: {result.get('language', 'unknown')}\n\n"

        tags = list(extra_tags or [])
        tags.extend(["youtube", "video_transcript", f"vid_{video_id}"])

        meta = {
            "source_url": url,
            "source_type": "youtube",
            "video_id": video_id,
            "video_title": title,
            "duration_seconds": result["duration_seconds"],
            "is_generated": result["is_generated"],
        }
        if extra_metadata:
            meta.update(extra_metadata)

        chunks = self.chunk_text(
            text=header + text,
            source_filename=title,
            extra_tags=tags,
            extra_metadata=meta,
        )

        return {
            "chunks": chunks,
            "title": title,
            "video_id": video_id,
            "duration_seconds": result["duration_seconds"],
            "transcript_length": len(text),
            "is_generated": result["is_generated"],
            "language": result.get("language", "unknown"),
        }

    # ─────────────────────────────────────────────
    # TEXT EXTRACTION (per format)
    # ─────────────────────────────────────────────

    def _extract_text(self, file_obj: BinaryIO, ext: str, filename: str) -> str:
        """Route to the correct extractor based on file extension."""
        extractors = {
            ".pdf": self._extract_pdf,
            ".docx": self._extract_docx,
            ".txt": self._extract_plaintext,
            ".md": self._extract_plaintext,
            ".csv": self._extract_csv,
            ".json": self._extract_json,
        }
        extractor = extractors.get(ext)
        if extractor is None:
            raise ValueError(f"No extractor for: {ext}")

        text = extractor(file_obj, filename)
        return self._clean_text(text)

    def _extract_pdf(self, file_obj: BinaryIO, filename: str) -> str:
        """Extract text from PDF files with smart Gemini Vision strategy.

        Strategy:
        1. Small PDFs (< 30 pages) + Gemini available: Vision for everything
        2. Large PDFs + Gemini available:
           a. Detect which pages have images/formulas/figures
           b. If > 50% pages are rich → full Vision extraction in page-range batches
           c. If <= 50% → pypdf base text + Vision only for rich pages
        3. No Gemini key → pypdf text-only fallback

        Calls self._progress_callback(message) if set, for UI progress updates.
        """
        pdf_bytes = file_obj.read()
        self._last_pdf_method = "text"

        api_key = os.environ.get("GOOGLE_API_KEY")

        try:
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(pdf_bytes))
            page_count = len(reader.pages)
        except Exception:
            page_count = 0
            reader = None

        if not api_key:
            file_obj.seek(0)
            return self._extract_pdf_text(file_obj, filename)

        if page_count > 0 and page_count < 30:
            vision_text = self._extract_pdf_vision(pdf_bytes, filename, api_key)
            if vision_text and len(vision_text.strip()) > 50:
                self._last_pdf_method = "gemini_vision"
                return vision_text
            file_obj.seek(0)
            return self._extract_pdf_text(file_obj, filename)

        if page_count >= 30 and reader:
            self._emit_progress(f"Analyzing {page_count} pages for images, formulas, figures...")
            rich_pages = self._detect_rich_pages(reader)
            rich_ratio = len(rich_pages) / page_count if page_count > 0 else 0

            if rich_ratio > 0.5:
                self._emit_progress(
                    f"Book is {rich_ratio:.0%} rich content — using full Gemini Vision "
                    f"({len(rich_pages)}/{page_count} pages have images/formulas)"
                )
                result = self._extract_large_pdf_full_vision(
                    pdf_bytes, filename, api_key, reader, page_count
                )
                if result:
                    return result
            elif rich_pages:
                file_obj.seek(0)
                base_text = self._extract_pdf_text(file_obj, filename)
                self._emit_progress(
                    f"Enriching {len(rich_pages)} pages with Gemini Vision "
                    f"({page_count - len(rich_pages)} pages text-only)"
                )
                enriched = self._enrich_large_pdf_with_vision(
                    pdf_bytes, filename, api_key, reader, base_text
                )
                if enriched:
                    return enriched
                self._last_pdf_method = "text"
                return base_text
            else:
                self._emit_progress(f"No rich content detected — using text extraction")
                file_obj.seek(0)
                return self._extract_pdf_text(file_obj, filename)

        file_obj.seek(0)
        return self._extract_pdf_text(file_obj, filename)

    def _emit_progress(self, message: str):
        """Send a progress update if a callback is registered."""
        print(f"  ℹ {message}")
        cb = getattr(self, "_progress_callback", None)
        if cb:
            cb(message)

    def _extract_large_pdf_full_vision(
        self, pdf_bytes: bytes, filename: str, api_key: str,
        reader, page_count: int,
    ) -> Optional[str]:
        """Extract ALL pages through Gemini Vision in batches (for heavily rich PDFs).

        Splits the PDF into page-range batches and sends each batch to Gemini.
        Combines all results into a single text with page markers.
        """
        try:
            from pypdf import PdfWriter
            from google import genai
            from google.genai import types
            import io
        except ImportError:
            return None

        BATCH_SIZE = 20
        client = genai.Client(api_key=api_key)
        all_pages_text = [""] * page_count
        total_batches = (page_count + BATCH_SIZE - 1) // BATCH_SIZE

        print(f"  🔬 Full Vision extraction: {page_count} pages in {total_batches} batches of {BATCH_SIZE}")

        for batch_idx in range(total_batches):
            start = batch_idx * BATCH_SIZE
            end = min(start + BATCH_SIZE, page_count)
            batch_pages = list(range(start, end))

            writer = PdfWriter()
            for pg_idx in batch_pages:
                writer.add_page(reader.pages[pg_idx])

            buf = io.BytesIO()
            writer.write(buf)
            batch_bytes = buf.getvalue()

            page_range = f"{start + 1}-{end}"
            self._emit_progress(
                f"Gemini Vision batch {batch_idx + 1}/{total_batches}: pages {page_range} "
                f"({len(batch_bytes) // 1024} KB)"
            )

            page_list = ", ".join(str(p + 1) for p in batch_pages)
            prompt = (
                f"Extract ALL content from these PDF pages (pages {page_list}). Include:\n"
                "1. **All text** — every paragraph, heading, caption exactly as written\n"
                "2. **All formulas/equations** in LaTeX ($...$ inline, $$...$$ display)\n"
                "3. **All figures/charts/diagrams** — describe in [Figure N: ...] blocks\n"
                "4. **All tables** — as readable text tables\n"
                "Mark each page with '--- Page N ---' using the page numbers above.\n"
                "Do NOT skip or summarize. Extract EVERYTHING from every page."
            )

            try:
                if len(batch_bytes) < 20_000_000:
                    pdf_part = types.Part.from_bytes(data=batch_bytes, mime_type="application/pdf")
                else:
                    import tempfile
                    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                        tmp.write(batch_bytes)
                        tmp_path = tmp.name
                    uploaded = client.files.upload(file=tmp_path)
                    import time as _time
                    while uploaded.state.name == "PROCESSING":
                        _time.sleep(2)
                        uploaded = client.files.get(name=uploaded.name)
                    pdf_part = uploaded
                    os.unlink(tmp_path)

                response = client.models.generate_content(
                    model=self.vision_model,
                    contents=[pdf_part, prompt],
                )
                vision_text = response.text if hasattr(response, "text") else str(response)

                if vision_text and len(vision_text.strip()) > 20:
                    page_map = self._parse_vision_pages(vision_text)
                    if page_map:
                        for pg_num, content in page_map.items():
                            if 0 <= pg_num < page_count:
                                all_pages_text[pg_num] = content
                        print(f"    ✓ Extracted {len(page_map)} pages")
                    else:
                        for pg_idx in batch_pages:
                            all_pages_text[pg_idx] = vision_text if pg_idx == batch_pages[0] else ""
                        print(f"    ✓ Extracted batch (couldn't split into pages)")

            except Exception as e:
                print(f"    ⚠ Batch {batch_idx + 1} failed ({e}), falling back to pypdf for these pages")
                for pg_idx in batch_pages:
                    text = (reader.pages[pg_idx].extract_text() or "").strip()
                    if text:
                        all_pages_text[pg_idx] = self._normalize_pdf_whitespace(text)

        self._last_pdf_method = "gemini_vision"
        filled = sum(1 for t in all_pages_text if t.strip())
        print(f"  ✓ Full Vision complete: {filled}/{page_count} pages extracted")

        combined = []
        for i, text in enumerate(all_pages_text):
            if text.strip():
                combined.append(f"--- Page {i + 1} ---\n{text}")
        return "\n\n".join(combined)

    def _parse_vision_pages(self, vision_text: str) -> dict:
        """Parse Gemini Vision output into {page_number: content} dict."""
        page_map = {}
        for section in vision_text.split("--- Page "):
            section = section.strip()
            if not section:
                continue
            try:
                pg_num_str = section.split("---")[0].strip()
                pg_num = int(pg_num_str) - 1
                pg_content = "---".join(section.split("---")[1:]).strip()
                if pg_content:
                    page_map[pg_num] = pg_content
            except (ValueError, IndexError):
                pass
        return page_map

    def _detect_rich_pages(self, reader) -> List[dict]:
        """Analyze each PDF page to detect images, formulas, and figures.

        Returns a list of dicts: [{"page": int, "reasons": [...]}]
        Only pages that need Gemini Vision are included.
        """
        import re
        rich_pages = []

        GREEK = set("αβγδεζηθικλμνξοπρστυφχψωΑΒΓΔΕΖΗΘΙΚΛΜΝΞΟΠΡΣΤΥΦΧΨΩ")
        MATH_SYMBOLS = set("∫∑∏∂∇√∞±≈≠≤≥∈∉⊂⊃∪∩∧∨¬→←↔⇒⇐⇔∀∃")
        GARBLED_MATH_RE = re.compile(
            r'(?:'
            r'[a-zA-Z]\s*[=<>≈≠≤≥]\s*[a-zA-Z0-9]'  # single-char equations like "F = m a"
            r'|[²³⁴⁵⁶⁷⁸⁹₀₁₂₃₄₅₆₇₈₉]'             # unicode super/subscripts
            r'|\b[A-Z]\s*\(\s*[a-z]\s*\)'             # function notation like "F ( x )"
            r')',
        )
        FIGURE_REF_RE = re.compile(
            r'(?:Fig(?:ure|\.)\s*\d|Table\s+\d|Equation\s+\d|Eq\.\s*\(?\d)',
            re.IGNORECASE,
        )

        for i, page in enumerate(reader.pages):
            reasons = []

            # 1. Detect embedded images via PDF page resources
            try:
                resources = page.get("/Resources")
                if resources:
                    xobjects = resources.get("/XObject")
                    if xobjects:
                        xobj = xobjects.get_object() if hasattr(xobjects, "get_object") else xobjects
                        if isinstance(xobj, dict):
                            image_count = 0
                            for key, val in xobj.items():
                                try:
                                    obj = val.get_object() if hasattr(val, "get_object") else val
                                    subtype = obj.get("/Subtype", "")
                                    if subtype == "/Image":
                                        image_count += 1
                                except Exception:
                                    pass
                            if image_count > 0:
                                reasons.append(f"{image_count} embedded image(s)")
            except Exception:
                pass

            # 2. Detect garbled math / formula artifacts in extracted text
            text = (page.extract_text() or "").strip()
            if text:
                greek_count = sum(1 for c in text if c in GREEK)
                math_sym_count = sum(1 for c in text if c in MATH_SYMBOLS)
                garbled_matches = len(GARBLED_MATH_RE.findall(text))
                fig_refs = len(FIGURE_REF_RE.findall(text))

                if greek_count > 5 or math_sym_count > 2:
                    reasons.append(f"math symbols ({greek_count} greek, {math_sym_count} math)")
                if garbled_matches > 3:
                    reasons.append(f"{garbled_matches} garbled equation patterns")
                if fig_refs > 0:
                    reasons.append(f"{fig_refs} figure/table/equation ref(s)")

            # 3. Sparse text — likely a full-page figure or diagram
            if len(text) < 80:
                reasons.append("very sparse text (likely full-page figure)")

            if reasons:
                rich_pages.append({"page": i, "reasons": reasons})

        return rich_pages

    def _enrich_large_pdf_with_vision(
        self, pdf_bytes: bytes, filename: str, api_key: str,
        reader, base_text: str,
    ) -> Optional[str]:
        """For large PDFs with <= 50% rich pages: pypdf base + Vision for rich pages only."""
        try:
            from pypdf import PdfWriter
            from google import genai
            from google.genai import types
            import io
        except ImportError:
            return None

        pages_text = []
        for page in reader.pages:
            pages_text.append((page.extract_text() or "").strip())

        rich_pages = self._detect_rich_pages(reader)
        if not rich_pages:
            self._last_pdf_method = "text"
            return None

        rich_indices = [rp["page"] for rp in rich_pages]
        total_batches = (len(rich_indices) + 14) // 15

        BATCH_SIZE = 15
        client = genai.Client(api_key=api_key)
        enriched_count = 0

        for batch_idx, batch_start in enumerate(range(0, len(rich_indices), BATCH_SIZE)):
            batch = rich_indices[batch_start:batch_start + BATCH_SIZE]

            self._emit_progress(
                f"Vision enriching batch {batch_idx + 1}/{total_batches}: "
                f"{len(batch)} pages ({len(rich_indices) - batch_start - len(batch)} remaining)"
            )

            page_map = self._vision_extract_pages(client, reader, batch)
            if page_map:
                for pg_idx, content in page_map.items():
                    if pg_idx < len(pages_text):
                        pages_text[pg_idx] = content
                        enriched_count += 1

        if enriched_count > 0:
            self._last_pdf_method = "gemini_vision"
            print(f"  ✓ Hybrid: {len(reader.pages)} pages text + {enriched_count} vision-enriched")
        else:
            self._last_pdf_method = "text"

        combined = []
        for i, text in enumerate(pages_text):
            if text.strip():
                text = self._normalize_pdf_whitespace(text)
                combined.append(f"--- Page {i + 1} ---\n{text}")
        return "\n\n".join(combined)

    def _vision_extract_pages(self, client, reader, page_indices: list) -> dict:
        """Send a batch of PDF pages to Gemini Vision. Returns {page_num: content}."""
        try:
            from pypdf import PdfWriter
            from google.genai import types
            import io
        except ImportError:
            return {}

        writer = PdfWriter()
        for pg_idx in page_indices:
            writer.add_page(reader.pages[pg_idx])

        buf = io.BytesIO()
        writer.write(buf)
        batch_bytes = buf.getvalue()

        page_list = ", ".join(str(p + 1) for p in page_indices)

        prompt = (
            f"Extract ALL content from these PDF pages (pages {page_list}). Include:\n"
            "1. **All text** exactly as written\n"
            "2. **All formulas/equations** in LaTeX ($...$ inline, $$...$$ display)\n"
            "3. **All figures/charts/diagrams** described in [Figure N: ...] blocks\n"
            "4. **All tables** as readable text\n"
            "Mark each page with '--- Page N ---'. Extract EVERYTHING."
        )

        try:
            if len(batch_bytes) < 20_000_000:
                pdf_part = types.Part.from_bytes(data=batch_bytes, mime_type="application/pdf")
            else:
                import tempfile
                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                    tmp.write(batch_bytes)
                    tmp_path = tmp.name
                uploaded = client.files.upload(file=tmp_path)
                import time as _time
                while uploaded.state.name == "PROCESSING":
                    _time.sleep(2)
                    uploaded = client.files.get(name=uploaded.name)
                pdf_part = uploaded
                os.unlink(tmp_path)

            response = client.models.generate_content(
                model=self.vision_model,
                contents=[pdf_part, prompt],
            )
            vision_text = response.text if hasattr(response, "text") else str(response)

            if vision_text and len(vision_text.strip()) > 20:
                page_map = self._parse_vision_pages(vision_text)
                if page_map:
                    print(f"    ✓ Vision extracted {len(page_map)} pages")
                    return page_map
                elif len(page_indices) == 1:
                    return {page_indices[0]: vision_text}

        except Exception as e:
            print(f"    ⚠ Vision batch failed ({e})")

        return {}

    def _extract_pdf_text(self, file_obj: BinaryIO, filename: str) -> str:
        """Text-only PDF extraction via pypdf."""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError(
                "pypdf is required for PDF support. "
                "Install with: pip install pypdf"
            )

        reader = PdfReader(file_obj)
        pages = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                page_text = self._normalize_pdf_whitespace(page_text)
                pages.append(f"--- Page {i + 1} ---\n{page_text}")
        return "\n\n".join(pages)

    def _extract_pdf_vision(
        self, pdf_bytes: bytes, filename: str, api_key: str
    ) -> Optional[str]:
        """Extract PDF content using Gemini vision — captures formulas, images, diagrams.

        Uses inline bytes for PDFs < 20 MB, Gemini File API for larger files (up to 2 GB).
        """
        try:
            from google import genai
            from google.genai import types
        except ImportError:
            return None

        size_mb = len(pdf_bytes) / (1024 * 1024)
        print(f"  🔬 Using Gemini vision to extract '{filename}' ({size_mb:.1f} MB, captures formulas + images)...")

        client = genai.Client(api_key=api_key)

        prompt = (
            "Extract ALL content from this PDF document thoroughly. You must include:\n\n"
            "1. **All text content** — preserve every paragraph, heading, and caption exactly.\n"
            "2. **All formulas and equations** — write them in LaTeX notation wrapped in $...$ "
            "(inline) or $$...$$ (display). Convert every symbol, fraction, subscript, "
            "and superscript faithfully.\n"
            "3. **All figures, charts, and diagrams** — describe each one in detail inside a "
            "[Figure N: ...] block. Include axis labels, data trends, legends, and key values "
            "visible in the figure.\n"
            "4. **All tables** — reproduce them as readable text tables with aligned columns.\n"
            "5. **Section headings** — preserve the document structure with heading levels.\n\n"
            "Output clean, readable text organized by page. "
            "Mark page boundaries with '--- Page N ---'.\n"
            "Do NOT skip or summarize any content. Extract EVERYTHING."
        )

        try:
            INLINE_LIMIT = 20_000_000  # 20 MB

            if len(pdf_bytes) < INLINE_LIMIT:
                pdf_part = types.Part.from_bytes(data=pdf_bytes, mime_type="application/pdf")
            else:
                import tempfile, time
                print(f"  📤 File > 20 MB — uploading via Gemini File API...")
                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                    tmp.write(pdf_bytes)
                    tmp_path = tmp.name

                uploaded = client.files.upload(file=tmp_path)
                while uploaded.state.name == "PROCESSING":
                    time.sleep(2)
                    uploaded = client.files.get(name=uploaded.name)

                if uploaded.state.name == "FAILED":
                    print(f"  ⚠ Gemini File API upload failed")
                    os.unlink(tmp_path)
                    return None

                print(f"  ✓ File uploaded to Gemini ({uploaded.name})")
                pdf_part = uploaded
                os.unlink(tmp_path)

            response = client.models.generate_content(
                model=self.vision_model,
                contents=[pdf_part, prompt],
            )
            text = response.text if hasattr(response, "text") else str(response)

            if text and len(text.strip()) > 50:
                print(f"  ✓ Gemini vision extracted {len(text):,} chars from '{filename}'")
                return text
            else:
                print(f"  ⚠ Gemini vision returned insufficient content, falling back to text extraction")
                return None

        except Exception as e:
            print(f"  ⚠ Gemini vision failed ({e}), falling back to text extraction")
            return None

    @staticmethod
    def _normalize_pdf_whitespace(text: str) -> str:
        """Clean up PDF extraction artifacts where every word lands on its own line."""
        paragraphs = re.split(r"\n{2,}", text)
        tokens: list[str] = []
        for para in paragraphs:
            stripped = para.strip()
            if stripped:
                tokens.append(re.sub(r"\s+", " ", stripped))
        if not tokens:
            return text

        result: list[str] = [tokens[0]]
        for tok in tokens[1:]:
            prev = result[-1]
            if len(tok.split()) <= 3 or not prev.rstrip()[-1:] in ".!?":
                result[-1] = prev.rstrip() + " " + tok
            else:
                result.append(tok)
        return "\n\n".join(result)

    def _extract_docx(self, file_obj: BinaryIO, filename: str) -> str:
        """Extract text from Word documents."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX support. "
                "Install with: pip install python-docx"
            )

        doc = Document(file_obj)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve heading hierarchy
                if para.style and para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading", "").strip()
                    prefix = "#" * int(level) if level.isdigit() else "#"
                    paragraphs.append(f"{prefix} {text}")
                else:
                    paragraphs.append(text)

        # Also extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                paragraphs.append("\n".join(rows))

        return "\n\n".join(paragraphs)

    def _extract_plaintext(self, file_obj: BinaryIO, filename: str) -> str:
        """Extract text from TXT/MD files."""
        raw = file_obj.read()
        # Try UTF-8 first, fall back to latin-1
        try:
            return raw.decode("utf-8")
        except UnicodeDecodeError:
            return raw.decode("latin-1", errors="replace")

    def _extract_csv(self, file_obj: BinaryIO, filename: str) -> str:
        """Extract text from CSV files, converting rows to readable text.

        Each row becomes a self-contained paragraph separated by blank lines
        so the chunker treats rows as atomic units and avoids splitting a
        person's name from their attributes.
        """
        raw = file_obj.read()
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            text = raw.decode("latin-1", errors="replace")

        reader = csv.reader(io.StringIO(text))
        rows = list(reader)
        if not rows:
            return ""

        headers = rows[0]
        paragraphs = []
        for row in rows[1:]:
            if not any(cell.strip() for cell in row):
                continue
            pairs = []
            for h, v in zip(headers, row):
                if v.strip():
                    pairs.append(f"{h.strip()}: {v.strip()}")
            if pairs:
                paragraphs.append(". ".join(pairs))

        return "\n\n".join(paragraphs)

    def _extract_json(self, file_obj: BinaryIO, filename: str) -> str:
        """Extract text from JSON files, flattening to readable text."""
        raw = file_obj.read()
        try:
            data = json.loads(raw.decode("utf-8"))
        except (UnicodeDecodeError, json.JSONDecodeError) as e:
            raise ValueError(f"Invalid JSON file: {e}")

        return self._flatten_json(data)

    # ─────────────────────────────────────────────
    # CHUNKING HELPERS
    # ─────────────────────────────────────────────

    def _split_paragraphs(self, text: str) -> List[str]:
        """Split text into paragraphs, respecting semantic boundaries."""
        # Split on double newlines (paragraph breaks)
        raw = re.split(r"\n\s*\n", text)
        paragraphs = [p.strip() for p in raw if p.strip()]
        return paragraphs

    def _merge_and_split(self, paragraphs: List[str]) -> List[str]:
        """
        Merge small paragraphs into chunks up to chunk_size,
        and split large paragraphs that exceed chunk_size.
        """
        chunks = []
        current_chunk = ""

        for para in paragraphs:
            # If a single paragraph exceeds chunk size, split it
            if len(para) > self.chunk_size:
                # Flush current chunk first
                if current_chunk:
                    chunks.append(current_chunk)
                    current_chunk = ""
                # Split the large paragraph by sentences
                sub_chunks = self._split_large_paragraph(para)
                chunks.extend(sub_chunks)
                continue

            # Would adding this paragraph exceed the chunk size?
            test = current_chunk + "\n\n" + para if current_chunk else para
            if len(test) > self.chunk_size:
                # Flush current chunk
                if current_chunk:
                    chunks.append(current_chunk)
                current_chunk = para
            else:
                current_chunk = test

        # Don't forget the last chunk
        if current_chunk:
            chunks.append(current_chunk)

        return chunks

    def _split_large_paragraph(self, text: str) -> List[str]:
        """Split a large paragraph into sentence-boundary chunks."""
        # Split by sentences
        sentences = re.split(r"(?<=[.!?])\s+", text)
        chunks = []
        current = ""

        for sentence in sentences:
            test = current + " " + sentence if current else sentence
            if len(test) > self.chunk_size and current:
                chunks.append(current)
                current = sentence
            else:
                current = test

        if current:
            chunks.append(current)

        return chunks

    def _detect_section_header(self, text: str) -> Optional[str]:
        """Try to detect a section header at the start of a chunk."""
        lines = text.split("\n", 3)
        for line in lines[:2]:
            line = line.strip()
            # Markdown headers
            if line.startswith("#"):
                return line.lstrip("#").strip()[:80]
            # ALL CAPS headers
            if line.isupper() and 3 < len(line) < 60:
                return line.title()
            # Numbered section (e.g., "1.2 Installation")
            if re.match(r"^\d+\.[\d.]*\s+\w", line):
                return line[:80]
        return None

    # ─────────────────────────────────────────────
    # UTILITY
    # ─────────────────────────────────────────────

    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean extracted text: normalize whitespace, remove control chars."""
        # Strip any leftover HTML tags (common in web page extraction)
        text = re.sub(r"<[^>]+>", " ", text)
        # Remove HTML entities
        text = re.sub(r"&[a-zA-Z]+;", " ", text)
        text = re.sub(r"&#\d+;", " ", text)
        # Remove null bytes and other control characters (keep newlines/tabs)
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
        # Normalize excessive whitespace within lines
        text = re.sub(r"[^\S\n]+", " ", text)
        # Normalize excessive blank lines
        text = re.sub(r"\n{4,}", "\n\n\n", text)
        return text.strip()

    @staticmethod
    def _flatten_json(data: Any, prefix: str = "", depth: int = 0) -> str:
        """Flatten a JSON structure into readable text lines."""
        if depth > 10:
            return str(data)[:200]

        lines = []

        if isinstance(data, dict):
            for key, value in data.items():
                full_key = f"{prefix}.{key}" if prefix else key
                if isinstance(value, (dict, list)):
                    lines.append(
                        DocumentIngestor._flatten_json(value, full_key, depth + 1)
                    )
                else:
                    lines.append(f"{full_key}: {value}")

        elif isinstance(data, list):
            for i, item in enumerate(data):
                item_prefix = f"{prefix}[{i}]" if prefix else f"[{i}]"
                if isinstance(item, (dict, list)):
                    lines.append(
                        DocumentIngestor._flatten_json(item, item_prefix, depth + 1)
                    )
                else:
                    lines.append(f"{item_prefix}: {item}")
        else:
            lines.append(f"{prefix}: {data}" if prefix else str(data))

        return "\n".join(lines)
